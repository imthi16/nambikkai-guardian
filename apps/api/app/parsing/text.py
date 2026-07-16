"""Plain-text, Markdown, and DOCX extraction (single logical page)."""

import io

import docx

from app.parsing.types import ParsedDocument, ParsedPage, ParserError


def parse_text(data: bytes) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as error:
        msg = "text content is not valid UTF-8"
        raise ParserError(msg) from error
    return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)], parser="text")


def parse_docx(data: bytes) -> ParsedDocument:
    """Practical DOCX support: paragraph text in order, one logical page."""
    try:
        document = docx.Document(io.BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
    except Exception as error:  # noqa: BLE001 - python-docx raises varied types
        msg = "the DOCX archive could not be parsed"
        raise ParserError(msg) from error
    return ParsedDocument(
        pages=[ParsedPage(page_number=1, text="\n".join(paragraphs).strip())],
        parser="docx",
    )
